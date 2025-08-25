# app.py
from flask import Flask, render_template, request, redirect, url_for, send_file, session, flash
import sqlite3
import io
import csv, re
from docx import Document
from fpdf import FPDF
import os
from werkzeug.utils import secure_filename
from datetime import datetime, date, timedelta
from flask import jsonify
from unicodedata import normalize
from werkzeug.security import generate_password_hash, check_password_hash
import unicodedata

app = Flask(__name__)
app.secret_key = 'tu_clave_secreta_super_segura'
app.permanent_session_lifetime = timedelta(days=30)

# Config de subida de imágenes
UPLOAD_FOLDER = 'static/img/herramientas'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


# -------------------- Utilidades BD --------------------
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_db():
    db_path = os.path.join(app.root_path, "asistencias.db")
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

def col(row, key, default=None):
    """Devuelve row[key] si existe y no es None; sino default."""
    try:
        if hasattr(row, "keys") and key in row.keys() and row[key] is not None:
            return row[key]
    except Exception:
        pass
    return default

def table_columns(conn, table):
    cur = conn.execute(f"PRAGMA table_info({table})")
    return {r[1] for r in cur.fetchall()}

def update_user_fields(conn, user_id, fields: dict):
    """Actualiza solo columnas existentes en usuarios."""
    cols = table_columns(conn, "usuarios")
    data = {k: v for k, v in fields.items() if k in cols}
    if not data:
        return
    sets = ", ".join([f"{k}=?" for k in data.keys()])
    sql = f"UPDATE usuarios SET {sets}, updated_at=CURRENT_TIMESTAMP WHERE id=?"
    params = list(data.values()) + [user_id]
    conn.execute(sql, params)
    conn.commit()

def insert_row(conn, table, data: dict):
    """Inserta solo las columnas que existan en la tabla."""
    cols = table_columns(conn, table)
    filt = {k: v for k, v in data.items() if k in cols}
    if not filt:
        return None
    placeholders = ", ".join(["?"] * len(filt))
    sql = f"INSERT INTO {table} ({', '.join(filt.keys())}) VALUES ({placeholders})"
    conn.execute(sql, list(filt.values()))
    conn.commit()
    try:
        rid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        return rid
    except Exception:
        return None


# -------------------- Importación CSV tolerante --------------------
def _norm_key(s: str) -> str:
    s = (s or "").strip()
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    s = s.lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    return s.strip("_")

_HEADER_MAP = {
    "id": "external_id",
    "external_id": "external_id",
    "nombre": "nombre",
    "referencia": "referencia",
    "barrio": "barrio",
    "telefono": "telefono",
    "telefono_1": "telefono",
    "tel": "telefono",
    "tefono": "telefono",
    "celular": "telefono",
    "situacion": "situacion",
    "estado": "situacion",
    "exonerado": "exonerado",
    "exento": "exonerado",
    "tipo_valor": "tipo_valor",     # p.ej. "cliente 130.000"
    "tipo": "tipo",                 # si ya viene separado
    "valor": "valor",
    "vencimiento": "vencimiento",
    "fecha_vencimiento": "vencimiento",
}

def _parse_bool(v):
    s = (str(v or "")).strip().lower()
    return 1 if s in ("1","si","sí","true","verdadero","x","s","y","yes") else 0

def _parse_date_to_iso(v):
    s = (str(v or "")).strip()
    if not s: return None
    m = re.match(r"^(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{2,4})$", s)
    if m:
        d, mth, y = m.groups()
        y = "20"+y if len(y)==2 else y
        try:
            return date(int(y), int(mth), int(d)).isoformat()
        except Exception:
            return s
    if re.match(r"^\d{4}-\d{2}-\d{2}$", s):
        return s
    return s

def _only_digits(s):
    return re.sub(r"\D+", "", str(s or ""))

def _try_decode(file_storage):
    data = file_storage.read()
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(enc), enc
        except Exception:
            continue
    return data.decode("latin-1", errors="ignore"), "latin-1"

def _guess_delimiter(text):
    try:
        dialect = csv.Sniffer().sniff(text[:1024], delimiters=[",",";","|","\t"])
        return dialect.delimiter
    except Exception:
        head = text.splitlines()[0] if text.splitlines() else ""
        return ";" if head.count(";") > head.count(",") else ","

def _split_tipo_valor(raw_tipo_o_tv, raw_valor):
    """
    - tipo: si viene vacío -> 'cliente'
    - valor: toma 'valor' si viene; si no, intenta extraer número de 'tipo_valor'
    """
    tipo = (raw_tipo_o_tv or "").strip().lower()
    if not tipo:
        tipo = "cliente"

    valor = (raw_valor or "").strip()
    if not valor and raw_tipo_o_tv:
        m = re.search(r"(\d[\d\.\,]*)", str(raw_tipo_o_tv))
        if m:
            valor = m.group(1).strip()
    return tipo, (valor or None)


# ===========================
#  AUTH: Registro / Login / Forgot
# ===========================
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        usuario = request.form.get("usuario","").strip()
        contrasena = request.form.get("contrasena","")
        email = request.form.get("email","").strip()
        nombre = request.form.get("nombre","").strip()

        if not usuario or not contrasena:
            flash("Usuario y contraseña son obligatorios", "danger")
            return render_template("register.html")

        db = get_db()
        if db.execute("SELECT 1 FROM usuarios WHERE usuario=?", (usuario,)).fetchone():
            db.close()
            flash("Ese usuario ya existe", "danger")
            return render_template("register.html")

        cols = table_columns(db, "usuarios")
        if "password_hash" in cols:
            pwd_hash = generate_password_hash(contrasena)
            db.execute("""
                INSERT INTO usuarios (usuario, password_hash, email, nombre, rol, foto_url, created_at, updated_at)
                VALUES (?, ?, ?, ?, 'operador', NULL, datetime('now'), datetime('now'))
            """, (usuario, pwd_hash, email or None, nombre or usuario))
        else:
            db.execute("INSERT INTO usuarios (usuario, contrasena) VALUES (?,?)", (usuario, contrasena))

        db.commit()
        db.close()
        flash("Usuario registrado. Ya podés iniciar sesión.", "success")
        return redirect(url_for("login"))

    return render_template("register.html")

@app.route("/forgot_password", methods=["GET","POST"])
def forgot_password():
    if request.method == "POST":
        usuario = request.form.get("usuario","").strip()
        nueva = request.form.get("nueva","")
        confirmar = request.form.get("confirmar","")
        if not usuario or not nueva or not confirmar:
            flash("Completá todos los campos","danger")
            return render_template("forgot_password.html")
        if nueva != confirmar:
            flash("Las contraseñas no coinciden","danger")
            return render_template("forgot_password.html")

        db = get_db()
        user = db.execute("SELECT * FROM usuarios WHERE usuario=?", (usuario,)).fetchone()
        if not user:
            db.close()
            flash("El usuario no existe","danger")
            return render_template("forgot_password.html")

        cols = table_columns(db, "usuarios")
        if "password_hash" in cols:
            pwd_hash = generate_password_hash(nueva)
            update_user_fields(db, col(user, "id"), {"password_hash": pwd_hash, "contrasena": None})
        else:
            update_user_fields(db, col(user, "id"), {"contrasena": nueva})

        flash("Contraseña actualizada. Iniciá sesión.","success")
        return redirect(url_for("login"))

    return render_template("forgot_password.html")

@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        usuario = request.form.get("usuario", "").strip()
        contrasena = request.form.get("contrasena", "")

        db = get_db()
        user = db.execute("SELECT * FROM usuarios WHERE usuario = ?", (usuario,)).fetchone()
        db.close()

        ok = False
        if user:
            pwd_hash = col(user, "password_hash")
            plano    = col(user, "contrasena")
            if pwd_hash:
                ok = check_password_hash(pwd_hash, contrasena)
            elif plano is not None:
                ok = (contrasena == plano)

        if user and ok:
            session.clear()
            session["usuario_id"] = col(user, "id")
            session["usuario"]    = col(user, "usuario")
            session["nombre"]     = col(user, "nombre", col(user, "usuario"))
            session["email"]      = col(user, "email")
            session["rol"]        = col(user, "rol", "operador")

            foto_rel = col(user, "foto_url")
            session["foto_url"] = url_for('static', filename=foto_rel) if foto_rel else None

            session.permanent = bool(request.form.get("recordarme"))
            return redirect(url_for("menu"))
        else:
            flash("Usuario o contraseña incorrectos", "danger")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ===========================
#  Páginas principales
# ===========================
@app.route("/menu")
def menu():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))
    return render_template("menu.html")


@app.route("/nuevo_ticket", methods=["GET", "POST"])
def nuevo_ticket():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    if request.method == "POST":
        # --- Campos del formulario ---
        cliente        = (request.form.get("cliente") or "").strip()
        cliente_id     = request.form.get("cliente_id") or None
        direccion      = (request.form.get("direccion") or "").strip()
        tipo           = request.form.get("tipo") or "Soporte"
        prioridad      = request.form.get("prioridad") or "Media"
        tecnico_nombre = (request.form.get("tecnico") or "").strip()
        tecnico_id     = request.form.get("tecnico_id") or None
        problema       = (request.form.get("problema") or "").strip()
        cedula         = (request.form.get("cedula") or "").strip()
        pppoe          = (request.form.get("pppoe") or "").strip()
        canal          = (request.form.get("canal") or "web").strip()
        estado         = (request.form.get("estado") or "pendiente").strip()

        # Fecha/hora agenda
        programada_local = (request.form.get("programada_local") or "").strip()
        programada_en = programada_local.replace("T", " ") if programada_local else None

        # PPPoE si falta
        if not pppoe and cliente:
            slug = normalize("NFD", cliente.lower()).encode("ascii", "ignore").decode("ascii")
            slug = "".join(ch for ch in slug if ch.isalnum())
            pppoe = f"{slug}@spynet.com"

        fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        db = get_db()
        insert_row(db, "asistencias", {
            "cliente": cliente,
            "direccion": direccion,
            "tipo": tipo,
            "prioridad": prioridad,
            "tecnico": tecnico_nombre,
            "problema": problema,
            "fecha": fecha,
            "pppoe": pppoe,
            "cliente_id": cliente_id,
            "cedula": cedula,
            "programada_en": programada_en,
            "estado": estado,
            "canal": canal,
            "tecnico_id": tecnico_id,
        })
        db.close()

        flash("Asistencia registrada.", "success")
        return redirect(url_for("tickets"))

    # --- GET: cargar selects ---
    db = get_db()
    ccols = table_columns(db, "clientes")

    # columnas base + opcionales (solo si existen en tu tabla)
    base_cols = ["id", "nombre", "apellido", "direccion"]
    opt_cols  = ["cedula", "pppoe", "telefono", "barrio", "referencia", "tipo_valor", "valor", "plan", "tipo"]

    select_cols = base_cols + [c for c in opt_cols if c in ccols]
    sql = f"SELECT {', '.join(select_cols)} FROM clientes WHERE activo=1 ORDER BY nombre COLLATE NOCASE ASC"
    clientes_rows = db.execute(sql).fetchall()

    clientes = []
    for r in clientes_rows:
        clientes.append({
            "id": r["id"],
            "nombre": r["nombre"],
            "apellido": col(r, "apellido", ""),
            "direccion": col(r, "direccion", ""),
            "cedula": col(r, "cedula", ""),
            "pppoe": col(r, "pppoe", ""),
            "telefono": col(r, "telefono", ""),
            "barrio": col(r, "barrio", ""),
            "referencia": col(r, "referencia", ""),
            "tipo_valor": col(r, "tipo_valor", ""),
            "valor": col(r, "valor", col(r, "plan", "")),  # usa valor; si no, plan
            "plan": col(r, "plan", ""),
            "tipo": col(r, "tipo", "cliente"),
        })

    tecnicos = db.execute("""
        SELECT id, nombre
        FROM tecnicos
        WHERE activo = 1
        ORDER BY nombre COLLATE NOCASE ASC
    """).fetchall()
    db.close()

    return render_template("nuevo_ticket.html", clientes=clientes, tecnicos=tecnicos)


@app.route("/tickets")
def tickets():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    db = get_db()
    rows = db.execute("SELECT * FROM asistencias ORDER BY datetime(fecha) DESC").fetchall()
    data = [dict(row) for row in rows]
    db.close()
    return render_template("tickets.html", tickets=data)


# ===========================
#  Descargas (PDF/WORD)
# ===========================
@app.route("/descargar/pdf")
def descargar_pdf():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    db = get_db()
    rows = db.execute("SELECT * FROM asistencias ORDER BY datetime(fecha) DESC").fetchall()
    tickets = [dict(row) for row in rows]
    db.close()

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Tickets de Asistencia", ln=True, align="C")

    for t in tickets:
        pdf.ln(4)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, txt=(
            f"Cliente: {t.get('cliente','')}\n"
            f"Dirección: {t.get('direccion','')}\n"
            f"Técnico: {t.get('tecnico','')}\n"
            f"Tipo: {t.get('tipo','')}\n"
            f"Prioridad: {t.get('prioridad','')}\n"
            f"PPPoE: {t.get('pppoe','')}\n"
            f"Problema: {t.get('problema') or 'N/A'}\n"
            f"Fecha: {t.get('fecha','')}"
        ))

    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    output = io.BytesIO(pdf_bytes)
    output.seek(0)
    return send_file(output, download_name="asistencias.pdf", as_attachment=True, mimetype="application/pdf")

@app.route("/descargar/word")
def descargar_word():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    db = get_db()
    rows = db.execute("SELECT * FROM asistencias ORDER BY datetime(fecha) DESC").fetchall()
    tickets = [dict(row) for row in rows]
    db.close()

    doc = Document()
    doc.add_heading("Tickets de Asistencia", 0)

    for t in tickets:
        doc.add_paragraph(f"Cliente: {t.get('cliente','')}")
        doc.add_paragraph(f"Dirección: {t.get('direccion','')}")
        doc.add_paragraph(f"Técnico: {t.get('tecnico','')}")
        doc.add_paragraph(f"Tipo: {t.get('tipo','')}")
        doc.add_paragraph(f"Prioridad: {t.get('prioridad','')}")
        doc.add_paragraph(f"PPPoE: {t.get('pppoe','')}")
        doc.add_paragraph(f"Problema: {t.get('problema') or 'N/A'}")
        doc.add_paragraph(f"Fecha: {t.get('fecha','')}")
        doc.add_paragraph("")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return send_file(output, download_name="asistencias.docx", as_attachment=True,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ===========================
#  Otras páginas
# ===========================
@app.route("/mapa")
def mapa():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))
    return render_template("mapa.html")

@app.route("/instalaciones")
def instalaciones():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))
    return render_template("instalaciones.html")

@app.route("/equipos")
def equipos():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    conn = get_db()
    equipos = conn.execute("SELECT * FROM equipos").fetchall()
    herramientas = conn.execute("SELECT * FROM herramientas").fetchall()
    historial = conn.execute("""
        SELECT u.id, u.item_type, 
               CASE 
                 WHEN u.item_type = 'herramienta' THEN h.nombre 
                 WHEN u.item_type = 'equipo' THEN e.nombre 
                 ELSE 'Desconocido' END AS nombre_item,
               u.tecnico, u.fecha, u.servicio
        FROM uso_items u
        LEFT JOIN herramientas h ON u.item_type = 'herramienta' AND u.item_id = h.id
        LEFT JOIN equipos e ON u.item_type = 'equipo' AND u.item_id = e.id
        ORDER BY datetime(u.fecha) DESC
        LIMIT 10
    """).fetchall()

    total_equipos = conn.execute("SELECT COUNT(*) FROM equipos").fetchone()[0]
    total_herramientas = conn.execute("SELECT COUNT(*) FROM herramientas").fetchone()[0]
    total_items = total_equipos + total_herramientas

    hoy_str = date.today().isoformat()
    en_uso_equipos = conn.execute("""
        SELECT COUNT(DISTINCT item_id) FROM uso_items 
        WHERE item_type = 'equipo' AND fecha LIKE ?
    """, (hoy_str + '%',)).fetchone()[0]

    en_uso_herramientas = conn.execute("""
        SELECT COUNT(DISTINCT item_id) FROM uso_items 
        WHERE item_type = 'herramienta' AND fecha LIKE ?
    """, (hoy_str + '%',)).fetchone()[0]

    en_uso = en_uso_equipos + en_uso_herramientas
    disponibles = total_items - en_uso if total_items >= en_uso else 0

    conn.close()

    now = datetime.now()
    return render_template("equipos.html", equipos=equipos, herramientas=herramientas,
                           historial=historial, total=total_items, en_uso=en_uso, disponibles=disponibles,
                           now=now)


@app.route('/registrar_equipo', methods=['POST'])
def registrar_equipo():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    nombre = request.form['nombre']
    tipo = request.form['tipo']
    descripcion = request.form.get('descripcion', '')

    conn = get_db()
    conn.execute("""
        INSERT INTO equipos (nombre, tipo, descripcion) VALUES (?, ?, ?)
    """, (nombre, tipo, descripcion))
    conn.commit()
    conn.close()

    flash("Equipo registrado correctamente.", "success")
    return redirect(url_for('equipos'))

@app.route('/registrar_uso_item', methods=['POST'])
def registrar_uso_item():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    item_id = request.form.get('item_id')
    tecnico = request.form.get('tecnico')
    servicio = request.form.get('servicio', '')
    fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    conn = get_db()
    cursor = conn.cursor()

    cursor.execute("SELECT id FROM equipos WHERE id = ?", (item_id,))
    es_equipo = cursor.fetchone() is not None

    cursor.execute("SELECT id FROM herramientas WHERE id = ?", (item_id,))
    es_herramienta = cursor.fetchone() is not None

    if es_equipo:
        item_type = 'equipo'
    elif es_herramienta:
        item_type = 'herramienta'
    else:
        flash("Elemento no encontrado", "danger")
        conn.close()
        return redirect(url_for('equipos'))

    cursor.execute("""
        INSERT INTO uso_items (item_type, item_id, tecnico, fecha, servicio)
        VALUES (?, ?, ?, ?, ?)
    """, (item_type, item_id, tecnico, fecha, servicio))
    conn.commit()
    conn.close()

    flash("Uso registrado correctamente.", "success")
    return redirect(url_for('equipos'))

@app.route('/subir_imagen_herramienta', methods=['POST'])
def subir_imagen_herramienta():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    if 'imagen' not in request.files:
        flash('No se ha seleccionado ningún archivo', 'danger')
        return redirect(url_for('equipos'))

    file = request.files['imagen']
    if file.filename == '':
        flash('No se ha seleccionado ningún archivo', 'danger')
        return redirect(url_for('equipos'))

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        herramienta_id = request.form.get('herramienta_id')
        if herramienta_id:
            conn = get_db()
            conn.execute('UPDATE herramientas SET imagen = ? WHERE id = ?', (filename, herramienta_id))
            conn.commit()
            conn.close()
            flash('Imagen subida correctamente', 'success')
        else:
            flash('No se indicó la herramienta', 'danger')
    else:
        flash('Formato de archivo no permitido', 'danger')

    return redirect(url_for('equipos'))

@app.route("/subir_foto_instalacion", methods=["POST"])
def subir_foto_instalacion():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    foto = request.files.get("foto")
    descripcion = request.form.get("descripcion", "")
    tecnico = session.get("usuario", "Desconocido")
    fecha_actual = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if not foto or foto.filename == "":
        flash("No se seleccionó ninguna foto", "danger")
        return redirect(url_for("equipos"))

    carpeta_destino = "static/img/herramientas"
    os.makedirs(carpeta_destino, exist_ok=True)

    nombre_archivo = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{secure_filename(foto.filename)}"
    ruta_guardado = os.path.join(carpeta_destino, nombre_archivo)
    foto.save(ruta_guardado)

    conn = get_db()
    conn.execute("""
        INSERT INTO fotos_asistencia (asistencia_id, tecnico, ruta_foto, descripcion, fecha)
        VALUES (?, ?, ?, ?, ?)
    """, (None, tecnico, ruta_guardado, descripcion, fecha_actual))
    conn.commit()
    conn.close()

    flash("Foto subida correctamente.", "success")
    return redirect(url_for("equipos"))


# ===========================
#  Clientes
# ===========================
@app.route("/clientes")
def clientes():
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))

    q        = request.args.get("q","").strip()
    situ     = request.args.get("situacion","").strip()   # filtro por situación
    exo      = request.args.get("exonerado","")           # "0" / "1" / ""
    barrio_f = request.args.get("barrio","").strip()

    db = get_db()
    # Traemos todo y normalizamos en Python para tolerar faltantes
    sql = "SELECT * FROM clientes WHERE 1=1"
    p = []
    if q:
        like = f"%{q}%"
        sql += " AND (IFNULL(nombre,'') LIKE ? OR IFNULL(telefono,'') LIKE ? OR IFNULL(referencia,'') LIKE ?)"
        p += [like, like, like]
    if situ:
        sql += " AND IFNULL(situacion,'') LIKE ?"
        p += [f"%{situ}%"]
    if exo in ("0","1"):
        sql += " AND exonerado=?"
        p += [int(exo)]
    if barrio_f:
        sql += " AND IFNULL(barrio,'') LIKE ?"
        p += [f"%{barrio_f}%"]
    sql += " ORDER BY nombre COLLATE NOCASE ASC"

    rows = db.execute(sql, p).fetchall()
    tot  = db.execute("SELECT COUNT(*) FROM clientes").fetchone()[0]
    db.close()

    # Normalizamos tipo/valor (fall back a tipo_valor si no existen)
    clientes_norm = []
    for r in rows:
        rd = dict(r)
        tipo = rd.get("tipo")
        valor = rd.get("valor")
        if not valor:
            tv = rd.get("tipo_valor") or ""
            m = re.search(r"(\d[\d\.\,]*)", tv)
            if m:
                valor = m.group(1)
        if not tipo:
            tipo = "cliente"
        rd["tipo"] = tipo
        rd["valor"] = valor
        clientes_norm.append(rd)

    return render_template("clientes.html",
                           clientes=clientes_norm,
                           q=q, situacion=situ, exonerado=exo, barrio=barrio_f, total=tot)

@app.route("/clientes/nuevo", methods=["GET","POST"])
def clientes_nuevo():
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))
    if request.method == "POST":
        external_id = request.form.get("external_id","").strip()
        nombre      = request.form.get("nombre","").strip()
        referencia  = request.form.get("referencia","").strip()
        barrio      = request.form.get("barrio","").strip()
        telefono    = request.form.get("telefono","").strip()
        situacion   = request.form.get("situacion","").strip()
        exonerado   = 1 if request.form.get("exonerado") == "on" else 0
        tipo        = request.form.get("tipo","").strip() or "cliente"
        valor       = request.form.get("valor","").strip() or None
        vencimiento = request.form.get("vencimiento","").strip()

        activo = 1 if any(x in situacion.lower() for x in ["activo","al día","al dia","en servicio","ok"]) else 0

        db = get_db()
        cols = table_columns(db, "clientes")
        data = {
            "external_id": external_id or None,
            "nombre": nombre,
            "referencia": referencia or None,
            "barrio": barrio or None,
            "telefono": telefono or None,
            "situacion": situacion or None,
            "exonerado": exonerado,
            "vencimiento": vencimiento or None,
            "activo": activo
        }
        if "tipo" in cols: data["tipo"] = tipo
        if "valor" in cols: data["valor"] = valor
        if "tipo_valor" in cols and "tipo" not in cols and "valor" not in cols:
            data["tipo_valor"] = f"{tipo} {valor}" if valor else tipo

        insert_row(db, "clientes", data)
        db.close()
        flash("Cliente creado.", "success")
        return redirect(url_for("clientes"))
    return render_template("cliente_form.html", mode="new", cliente=None)

@app.route("/clientes/<int:cid>")
def clientes_detalle(cid):
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))

    db = get_db()
    c = db.execute("SELECT * FROM clientes WHERE id=?", (cid,)).fetchone()
    if not c:
        db.close(); flash("Cliente no encontrado.", "warning")
        return redirect(url_for("clientes"))

    full_name = f"{c['nombre']} {c['apellido']}".strip() if "apellido" in c.keys() else c["nombre"]
    tickets = db.execute("""
        SELECT * FROM asistencias WHERE cliente = ?
        ORDER BY datetime(fecha) DESC LIMIT 10
    """, (full_name,)).fetchall()
    db.close()
    return render_template("cliente_detalle.html", c=c, tickets=tickets)

@app.route("/clientes/<int:cid>/editar", methods=["GET","POST"])
def clientes_editar(cid):
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))
    db = get_db()
    c = db.execute("SELECT * FROM clientes WHERE id=?", (cid,)).fetchone()
    if not c:
        db.close(); flash("Cliente no encontrado.", "warning"); return redirect(url_for("clientes"))

    if request.method == "POST":
        external_id = request.form.get("external_id","").strip()
        nombre      = request.form.get("nombre","").strip()
        referencia  = request.form.get("referencia","").strip()
        barrio      = request.form.get("barrio","").strip()
        telefono    = request.form.get("telefono","").strip()
        situacion   = request.form.get("situacion","").strip()
        exonerado   = 1 if request.form.get("exonerado") == "on" else 0
        tipo        = request.form.get("tipo","").strip() or "cliente"
        valor       = request.form.get("valor","").strip() or None
        vencimiento = request.form.get("vencimiento","").strip()
        activo      = 1 if any(x in situacion.lower() for x in ["activo","al día","al dia","en servicio","ok"]) else 0

        cols = table_columns(db, "clientes")
        data = {
            "external_id": external_id or None,
            "nombre": nombre,
            "referencia": referencia or None,
            "barrio": barrio or None,
            "telefono": telefono or None,
            "situacion": situacion or None,
            "exonerado": exonerado,
            "vencimiento": vencimiento or None,
            "activo": activo
        }
        if "tipo" in cols: data["tipo"] = tipo
        if "valor" in cols: data["valor"] = valor
        if "tipo_valor" in cols and "tipo" not in cols and "valor" not in cols:
            data["tipo_valor"] = f"{tipo} {valor}" if valor else tipo

        sets = ", ".join([f"{k}=?" for k in data.keys()])
        db.execute(f"UPDATE clientes SET {sets} WHERE id=?", list(data.values())+[cid])
        db.commit(); db.close()
        flash("Cliente actualizado.", "success")
        return redirect(url_for("clientes"))
    db.close()
    return render_template("cliente_form.html", mode="edit", cliente=c)


@app.route("/clientes/<int:cid>/toggle", methods=["POST"])
def clientes_toggle(cid):
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))
    db = get_db()
    cur = db.execute("SELECT activo FROM clientes WHERE id=?", (cid,)).fetchone()
    if cur:
        nuevo = 0 if cur["activo"]==1 else 1
        db.execute("UPDATE clientes SET activo=? WHERE id=?", (nuevo, cid))
        db.commit()
        flash("Estado actualizado.", "success")
    db.close()
    return redirect(url_for("clientes"))

@app.route("/clientes/<int:cid>/eliminar", methods=["POST"])
def clientes_eliminar(cid):
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))
    db = get_db()
    db.execute("DELETE FROM clientes WHERE id=?", (cid,))
    db.commit(); db.close()
    flash("Cliente eliminado.", "success")
    return redirect(url_for("clientes"))


# ===========================
#  Agenda + Acciones de tickets (POST)
# ===========================
@app.route("/agenda")
def agenda():
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    dia       = request.args.get("dia") or datetime.now().strftime("%Y-%m-%d")
    estado_f  = request.args.get("estado") or ""
    tecnico_f = request.args.get("tecnico_id") or ""

    db = get_db()

    try:
        tecnicos = db.execute(
            "SELECT id, nombre FROM tecnicos WHERE activo=1 ORDER BY nombre"
        ).fetchall()
    except sqlite3.OperationalError:
        tecnicos = []

    acols = table_columns(db, "asistencias")
    if "programada_en" not in acols:
        db.close()
        flash("Tu base de datos no tiene la columna 'programada_en' en asistencias. Ejecutá crear_db.py para actualizar.", "warning")
        return render_template("agenda.html",
                               eventos=[],
                               dia=dia,
                               prev_dia=(datetime.strptime(dia, "%Y-%m-%d") - timedelta(days=1)).strftime("%Y-%m-%d"),
                               next_dia=(datetime.strptime(dia, "%Y-%m-%d") + timedelta(days=1)).strftime("%Y-%m-%d"),
                               estado_f=estado_f, tecnico_f=tecnico_f, tecnicos=tecnicos)

    join_c = " LEFT JOIN clientes c ON a.cliente_id = c.id " if "cliente_id" in acols else " "
    join_t = " LEFT JOIN tecnicos t ON a.tecnico_id = t.id " if "tecnico_id" in acols else " "
    select_c = " , c.nombre AS c_nombre, c.apellido AS c_apellido " if "cliente_id" in acols else " , NULL AS c_nombre, NULL AS c_apellido "
    select_t = " , t.nombre AS t_nombre " if "tecnico_id" in acols else " , NULL AS t_nombre "

    sql = f"""
      SELECT a.* {select_c} {select_t}
        FROM asistencias a
        {join_c}
        {join_t}
       WHERE a.programada_en IS NOT NULL
         AND date(a.programada_en) = ?
    """
    params = [dia]
    if estado_f and "estado" in acols:
        sql += " AND a.estado = ?"
        params.append(estado_f)
    if tecnico_f and "tecnico_id" in acols:
        sql += " AND a.tecnico_id = ?"
        params.append(tecnico_f)

    sql += " ORDER BY time(a.programada_en) ASC"
    eventos = db.execute(sql, params).fetchall()
    db.close()

    d = datetime.strptime(dia, "%Y-%m-%d")
    prev_dia = (d - timedelta(days=1)).strftime("%Y-%m-%d")
    next_dia = (d + timedelta(days=1)).strftime("%Y-%m-%d")

    return render_template(
        "agenda.html",
        eventos=eventos,
        dia=dia,
        prev_dia=prev_dia,
        next_dia=next_dia,
        estado_f=estado_f,
        tecnico_f=tecnico_f,
        tecnicos=tecnicos,
    )


@app.route("/tickets/<int:tid>/programar", methods=["POST"])
def tickets_programar(tid):
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    prog = (request.form.get("programada_local") or "").strip()  # "YYYY-MM-DDTHH:MM"
    programada_en = prog.replace("T", " ") if prog else None

    db = get_db()
    acols = table_columns(db, "asistencias")
    if "programada_en" not in acols:
        db.close()
        flash("No existe la columna 'programada_en' en asistencias. Actualizá la BD.", "warning")
        return redirect(request.referrer or url_for("agenda"))

    db.execute("UPDATE asistencias SET programada_en=? WHERE id=?", (programada_en, tid))
    db.commit()
    db.close()

    flash("Cita reprogramada.", "success")
    return redirect(request.referrer or url_for("agenda"))


@app.route("/tickets/<int:tid>/estado", methods=["POST"])
def tickets_cambiar_estado(tid):
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    nuevo = (request.form.get("estado") or "").strip()
    db = get_db()
    acols = table_columns(db, "asistencias")
    if "estado" not in acols:
        db.close()
        flash("No existe la columna 'estado' en asistencias. Actualizá la BD.", "warning")
        return redirect(request.referrer or url_for("agenda"))

    if nuevo not in ("pendiente", "en_progreso", "resuelto", "cancelado"):
        db.close()
        flash("Estado inválido.", "warning")
        return redirect(request.referrer or url_for("agenda"))

    db.execute("UPDATE asistencias SET estado=? WHERE id=?", (nuevo, tid))
    db.commit()
    db.close()

    flash("Estado actualizado.", "success")
    return redirect(request.referrer or url_for("agenda"))


@app.route("/tickets/<int:tid>/asignar", methods=["POST"])
def tickets_asignar(tid):
    if "usuario" not in session and "usuario_id" not in session:
        return redirect(url_for("login"))

    tecnico_id = request.form.get("tecnico_id") or None

    db = get_db()
    acols = table_columns(db, "asistencias")
    if "tecnico_id" not in acols:
        db.close()
        flash("No existe la columna 'tecnico_id' en asistencias. Actualizá la BD.", "warning")
        return redirect(request.referrer or url_for("agenda"))

    db.execute("UPDATE asistencias SET tecnico_id=? WHERE id=?", (tecnico_id, tid))
    db.commit()
    db.close()

    flash("Técnico asignado.", "success")
    return redirect(request.referrer or url_for("agenda"))


# ===========================
#  Importar clientes (CSV)
# ===========================
@app.route("/clientes/importar", methods=["POST"])
def clientes_importar():
    if "usuario_id" not in session and "usuario" not in session:
        return redirect(url_for("login"))

    f = request.files.get("csvfile")  # el input del formulario debe llamarse csvfile
    if not f or f.filename == "":
        flash("Seleccioná un archivo CSV.", "warning")
        return redirect(url_for("clientes"))

    text, enc = _try_decode(f)
    delim = _guess_delimiter(text)
    reader = csv.DictReader(io.StringIO(text), delimiter=delim)
    if not reader.fieldnames:
        flash("No pude leer encabezados del CSV.", "danger")
        return redirect(url_for("clientes"))

    keymap, desconocidas = {}, []
    for h in reader.fieldnames:
        nk = _norm_key(h)
        tk = _HEADER_MAP.get(nk)
        if tk: keymap[h] = tk
        else: desconocidas.append(h)
    if not keymap:
        flash("No reconocí columnas del CSV. Revisá los encabezados.", "danger")
        return redirect(url_for("clientes"))
    if desconocidas:
        flash(f"Aviso: columnas ignoradas: {', '.join(desconocidas)}", "warning")

    db = get_db()
    cols = table_columns(db, "clientes")
    ins = upd = 0

    for raw in reader:
        if not any((str(v or "").strip() for v in raw.values())):
            continue

        norm = { keymap[k]: (raw.get(k) or "").strip() for k in raw if k in keymap }

        ext_id     = norm.get("external_id") or None
        nombre     = norm.get("nombre") or ""
        referencia = norm.get("referencia") or None
        barrio     = norm.get("barrio") or None
        telefono   = norm.get("telefono") or None
        if telefono:
            digits = _only_digits(telefono)
            telefono = digits if len(digits) >= 6 else telefono
        situacion  = norm.get("situacion") or ""
        exonerado  = _parse_bool(norm.get("exonerado"))
        tv         = norm.get("tipo_valor")  # si viene junto
        tipo_in    = norm.get("tipo")
        valor_in   = norm.get("valor")
        tipo_final, valor_final = _split_tipo_valor(tipo_in or tv, valor_in)
        venc       = _parse_date_to_iso(norm.get("vencimiento"))

        activo = 0 if situacion.lower() in ("inactivo","baja","suspendido","cancelado") else 1
        if not (ext_id or nombre or telefono):
            continue

        data = {
            "external_id": ext_id,
            "nombre": nombre,
            "referencia": referencia,
            "barrio": barrio,
            "telefono": telefono,
            "situacion": situacion or None,
            "exonerado": exonerado,
            "vencimiento": venc,
            "activo": activo
        }
        # Guardamos tipo/valor si existen; mantenemos tipo_valor para compatibilidad si es la única
        if "tipo" in cols:   data["tipo"] = tipo_final
        if "valor" in cols:  data["valor"] = valor_final
        if "tipo_valor" in cols and ("tipo" not in cols or "valor" not in cols):
            data["tipo_valor"] = tv or f"{tipo_final} {valor_final or ''}".strip()

        # upsert por external_id; luego por teléfono
        row = None
        if ext_id:
            row = db.execute("SELECT id FROM clientes WHERE external_id=?", (ext_id,)).fetchone()
        if not row and telefono:
            row = db.execute("SELECT id FROM clientes WHERE telefono=?", (telefono,)).fetchone()

        if row:
            sets = ", ".join([f"{k}=?" for k in data.keys()])
            db.execute(f"UPDATE clientes SET {sets} WHERE id=?", list(data.values())+[row["id"]])
            upd += 1
        else:
            qs = ", ".join(["?"]*len(data))
            db.execute(f"INSERT INTO clientes ({', '.join(data.keys())}) VALUES ({qs})", list(data.values()))
            ins += 1

    db.commit(); db.close()
    flash(f"Importación OK. Insertados {ins}, actualizados {upd}. (codificación {enc}, separador '{delim}')", "success")
    return redirect(url_for("clientes"))


    # --- API datos para el mapa ---
@app.route("/api/mapa_datos", endpoint="api_mapa_datos")
def api_mapa_datos():
    if "usuario" not in session and "usuario_id" not in session:
        return jsonify({"error":"no_auth"}), 401

    db = get_db()
    tickets = db.execute("""
        SELECT a.id, a.cliente, a.direccion, a.tipo, a.prioridad, a.estado,
               a.programada_en, a.lat, a.lng,
               COALESCE(t.nombre, a.tecnico) AS tecnico
        FROM asistencias a
        LEFT JOIN tecnicos t ON a.tecnico_id = t.id
        WHERE a.lat IS NOT NULL AND a.lng IS NOT NULL
          AND date(a.fecha) >= date('now','-15 day')
    """).fetchall()

    pos = db.execute("""
        SELECT tp.tecnico_id, tp.lat, tp.lng, tp.ts, te.nombre
        FROM tecnico_pos tp
        JOIN (
            SELECT tecnico_id, MAX(ts) AS mts
            FROM tecnico_pos GROUP BY tecnico_id
        ) x ON x.tecnico_id = tp.tecnico_id AND x.mts = tp.ts
        LEFT JOIN tecnicos te ON te.id = tp.tecnico_id
    """).fetchall()
    db.close()

    return jsonify({
        "tickets": [dict(r) for r in tickets],
        "tecnicos": [{"id": r["tecnico_id"], "nombre": r["nombre"],
                      "lat": r["lat"], "lng": r["lng"], "ts": r["ts"]} for r in pos]
    })


# --- Trayectoria de un técnico (por fecha) ---
@app.route("/api/tecnico_trayectoria/<int:tid>", endpoint="api_tecnico_trayectoria")
def api_tecnico_trayectoria(tid):
    if "usuario" not in session and "usuario_id" not in session:
        return jsonify({"error":"no_auth"}), 401

    desde = request.args.get("desde") or (datetime.now()-timedelta(days=1)).strftime("%Y-%m-%d")
    hasta = request.args.get("hasta") or datetime.now().strftime("%Y-%m-%d")

    db = get_db()
    rows = db.execute("""
        SELECT lat, lng, ts
        FROM tecnico_pos
        WHERE tecnico_id=? AND date(ts) BETWEEN ? AND ?
        ORDER BY ts ASC
    """, (tid, desde, hasta)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


# --- Ping GPS (móvil) ---
@app.route("/gps", methods=["GET","POST"], endpoint="gps_ping")
def gps_ping():
    tecnico_id = request.values.get("tecnico_id")
    lat = request.values.get("lat")
    lng = request.values.get("lng")
    if not tecnico_id or not lat or not lng:
        return "Faltan parametros (tecnico_id, lat, lng)", 400

    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    db = get_db()
    db.execute("INSERT INTO tecnico_pos (tecnico_id, lat, lng, ts) VALUES (?,?,?,?)",
               (tecnico_id, float(lat), float(lng), ts))

    cols = table_columns(db, "tecnicos")
    if {"lat","lng","pos_updated_at"} <= cols:
        db.execute("UPDATE tecnicos SET lat=?, lng=?, pos_updated_at=? WHERE id=?",
                   (float(lat), float(lng), ts, tecnico_id))
    db.commit(); db.close()
    return "ok"

    

    @app.get("/health")
def health():
    return {"ok": True}

@app.get("/")
def home():
    # Ruta mínima para verificar que Flask responde en Vercel
    return "Flask en Vercel ✅"



# ===========================
#  Main
# ===========================
if __name__ == "__main__":
    app.run(debug=True)